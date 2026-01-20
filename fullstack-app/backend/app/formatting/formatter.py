import os
import openai
from openai import AzureOpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime

def format_title_case(text):
    """Format text in proper title case, keeping small words lowercase except at the beginning"""
    small_words = {'of', 'and', 'for', 'the', 'in', 'on', 'at', 'to', 'a', 'an', 'as', 'but', 'or', 'nor', 'with', 'by', 'from'}
    words = text.split()
    formatted_words = []
    
    for i, word in enumerate(words):
        # Always capitalize the first word and words not in small_words list
        if i == 0 or word.lower() not in small_words:
            # Handle AI specifically
            if word.lower() == 'ai':
                formatted_words.append('AI')
            else:
                formatted_words.append(word.capitalize())
        else:
            formatted_words.append(word.lower())
    
    return ' '.join(formatted_words)

def convert_markdown_to_html(text):
    """Convert markdown formatting to HTML"""
    # Convert **text** to <strong>text</strong> - make sure to handle multiple occurrences
    text = re.sub(r'\*\*([^*]+?)\*\*', r'<strong>\1</strong>', text)
    return text

def remove_markdown_formatting(text):
    """Remove markdown formatting markers"""
    # Remove ** markers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    return text

def add_formatted_text_to_paragraph(paragraph, text):
    """Add text with markdown formatting to a Word paragraph"""
    # Split text by markdown bold patterns
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # This is bold text - remove ** and add as bold run
            bold_text = part[2:-2]  # Remove ** from both ends
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Regular text
            if part:  # Only add if not empty
                paragraph.add_run(part)

def add_content_with_bullets(doc, content):
    """Add content to document, handling bullet points properly"""
    # Check if content contains bullet points
    if '- ' in content and '\n- ' in content:
        # Split into lines and process
        lines = content.split('\n')
        current_paragraph_lines = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('- '):
                # First, add any accumulated non-bullet content
                if current_paragraph_lines:
                    para_text = '\n'.join(current_paragraph_lines).strip()
                    if para_text:
                        content_paragraph = doc.add_paragraph()
                        add_formatted_text_to_paragraph(content_paragraph, para_text)
                    current_paragraph_lines = []
                
                # Add bullet point
                bullet_text = line[2:].strip()  # Remove '- ' prefix
                bullet_paragraph = doc.add_paragraph(style='List Bullet')
                add_formatted_text_to_paragraph(bullet_paragraph, bullet_text)
            else:
                # Regular line, accumulate it
                if line:  # Only add non-empty lines
                    current_paragraph_lines.append(line)
        
        # Add any remaining non-bullet content
        if current_paragraph_lines:
            para_text = '\n'.join(current_paragraph_lines).strip()
            if para_text:
                content_paragraph = doc.add_paragraph()
                add_formatted_text_to_paragraph(content_paragraph, para_text)
    else:
        # No bullet points, handle as regular paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                content_paragraph = doc.add_paragraph()
                add_formatted_text_to_paragraph(content_paragraph, para_text.strip())

def format_report_as_html(report_content, company_data):
    """Convert the raw report to formatted HTML using a simple approach"""
    
    # Start with the title and basic HTML structure
    html = f"""
    <div style="font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; line-height: 1.6;">
        <h1 style="text-align: center; margin-bottom: 30px; color: #2c3e50;">AI ASSESSMENT AND CONSULTATION</h1>
        
        <!-- Company Info Section -->
        <div style="margin-bottom: 30px;">
            <p><strong>Company Name:</strong> {company_data['company_name']}</p>
            <p><strong>Country:</strong> {company_data['country']}</p>
            <p><strong>Consultation Date:</strong> {company_data['consultation_date']}</p>
            <p><strong>Expert(s):</strong> {company_data['experts']}</p>
            <p><strong>Customer manager:</strong> {company_data['customer_manager']}</p>
            <p><strong>Consultation Type:</strong> {company_data['consultation_type']}</p>
        </div>
    """
    
    # Find each section in the report
    section_patterns = [
        ("AI Maturity Level", r"\*\*AI Maturity Level:\*\*(.*?)(?=\*\*|\Z)"),
        ("Current Solution Development Stage", r"\*\*Current Solution Development Stage:\*\*(.*?)(?=\*\*|\Z)"),
        ("Validity of Concept and Authenticity of Problem Addressed", r"\*\*Validity of Concept and Authenticity of Problem Addressed:\*\*(.*?)(?=\*\*|\Z)"),
        ("Integration and Importance of AI in the Idea", r"\*\*Integration and Importance of AI.*?:\*\*(.*?)(?=\*\*|\Z)"),
        ("Identified Target Market and Customer Segments", r"\*\*Identified Target Market.*?:\*\*(.*?)(?=\*\*|\Z)"),
        ("Data Requirement Assessment", r"\*\*Data Requirement Assessment:\*\*(.*?)(?=\*\*|\Z)"),
        ("Data Collection Strategy", r"\*\*Data Collection Strategy:\*\*(.*?)(?=\*\*|\Z)"),
        ("Technical Expertise and Capability", r"\*\*Technical Expertise.*?:\*\*(.*?)(?=\*\*|\Z)"),
        ("Expectations from FAIR Services", r"\*\*Expectations from FAIR Services:\*\*(.*?)(?=\*\*|\Z)"),
        ("Recommendations", r"\*\*Recommendations:\*\*(.*?)(?=\*\*|\Z)")
        ]
    
    # Extract and add each section
    for section_title, pattern in section_patterns:
        matches = re.search(pattern, report_content, re.DOTALL | re.IGNORECASE)
        if matches:
            content = matches.group(1).strip()
            
            # Convert markdown formatting to HTML
            content = convert_markdown_to_html(content)
            
            # Format bullet points for any section that has them
            if "- " in content:
                # Split content into paragraphs and bullet sections
                formatted_content = ""
                lines = content.split('\n')
                current_paragraph = []
                bullet_items = []
                in_bullets = False
                
                for line in lines:
                    line = line.strip()
                    if line.startswith('- '):
                        # Starting bullet section
                        if current_paragraph:
                            # Add accumulated paragraph
                            para_text = ' '.join(current_paragraph).strip()
                            if para_text:
                                formatted_content += f"<p>{para_text}</p>"
                            current_paragraph = []
                        
                        # Add bullet item
                        bullet_text = line[2:].strip()  # Remove '- '
                        bullet_items.append(bullet_text)
                        in_bullets = True
                    elif line == "" and in_bullets:
                        # Empty line after bullets - end bullet section
                        if bullet_items:
                            formatted_content += "<ul>"
                            for item in bullet_items:
                                formatted_content += f"<li>{item}</li>"
                            formatted_content += "</ul>"
                            bullet_items = []
                        in_bullets = False
                    elif line == "":
                        # Empty line - paragraph break
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph).strip()
                            if para_text:
                                formatted_content += f"<p>{para_text}</p>"
                            current_paragraph = []
                    else:
                        # Regular text line
                        if in_bullets:
                            # This might be continuation of bullet item
                            if bullet_items:
                                bullet_items[-1] += " " + line
                        else:
                            current_paragraph.append(line)
                
                # Handle remaining content
                if bullet_items:
                    formatted_content += "<ul>"
                    for item in bullet_items:
                        formatted_content += f"<li>{item}</li>"
                    formatted_content += "</ul>"
                
                if current_paragraph:
                    para_text = ' '.join(current_paragraph).strip()
                    if para_text:
                        formatted_content += f"<p>{para_text}</p>"
                
                content = formatted_content if formatted_content else f"<p>{content}</p>"
            else:
                # No bullet points - handle as regular paragraphs
                content = content.replace("\n\n", "</p><p>")
                content = f"<p>{content}</p>"
            
            # Format section title for HTML
            formatted_title = format_title_case(section_title)
            html += f"""
            <div style="margin-bottom: 15px;">
                <h3 style="color: #2c3e50; border-bottom: 1px solid #eee;">{formatted_title}</h3>
                {content}
            </div>
            """
    
    # Add AI Maturity Levels section with smaller font
    html += """
    <hr style="margin: 30px 0; border: 0; border-top: 1px solid #eee;">
    
    <div style="font-size: 0.8em; color: #666;">
        <h4 style="color: #2c3e50; font-size: 0.9em;">AI Maturity Levels:</h4>
        
        <p><strong>Low:</strong> Companies in early stages of AI integration or development, typically in ideation phase with limited data, resources, expertise, and minimal AI understanding. AI is minimally or not used in workflows, with no data management processes or AI roadmap.</p>
        
        <p><strong>Moderate:</strong> Companies progressing in their AI journey beyond proof of concept with functional solutions. They have adequate data, resources, expertise, and AI understanding. AI is partially or fully integrated into workflows with established or developing data management processes and an AI roadmap.</p>
        
        <p><strong>High:</strong> Companies with advanced AI products and established customer base. AI is integrated into workflows with established data management processes and AI roadmap. They require assistance with specific technical details or developing new AI applications.</p>
    </div>
    </div>
    """
    
    return html

def create_word_doc(report_content, file_path, company_data):
    """Create a Word document from the report content using a non-tabular format"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Add title
    title = doc.add_heading('AI ASSESSMENT AND CONSULTATION', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set title text to black
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add company info
    company_info = [
        f"Company Name: {company_data['company_name']}",
        f"Country: {company_data['country']}",
        f"Consultation Date: {company_data['consultation_date']}",
        f"Expert(s): {company_data['experts']}",
        f"Customer manager: {company_data['customer_manager']}",
        f"Consultation Type: {company_data['consultation_type']}"
    ]
    
    for info in company_info:
        doc.add_paragraph(info)
    
    doc.add_paragraph()  # Add spacing
    
    # Extract sections from the report content
    sections = [
            "AI Maturity Level",
            "Current Solution Development Stage",
            "Validity of Concept and Authenticity of Problem Addressed",
            "Integration and Importance of AI in the Idea",
            "Identified Target Market and Customer Segments",
            "Data Requirement Assessment",
            "Data Collection Strategy",
            "Technical Expertise and Capability", 
            "Expectations from FAIR Services",
            "Recommendations"
            ]
    
    # Try to extract content for each section
    for section in sections:
        section_title = section.strip()
        
        # Format the section title correctly (proper title case)
        formatted_section_title = format_title_case(section_title)
        
        # Find section in the report content with improved pattern matching
        pattern = rf"(?i)\*\*\s*{re.escape(section)}:?\*\*\s*(.*?)(?=\n\s*\*\*|\n\s*-{10,}|\Z)"
        matches = re.findall(pattern, report_content, re.DOTALL | re.IGNORECASE | re.MULTILINE)
        
        content = ""
        if matches:
            # Handle the case where matches is a list
            if isinstance(matches, list) and len(matches) > 0:
                content = matches[0]
                # Handle if content is a tuple
                if isinstance(content, tuple) and len(content) > 0:
                    content = content[0]
        
        # If we didn't find content with the pattern above, try an alternative approach
        if not content:
            pattern = rf"(?i)\b{re.escape(section)}\s*:\s*(.*?)(?=\n\s*\*\*|\n\s*[A-Z]|\n\s*-{10,}|\Z)"
            matches = re.findall(pattern, report_content, re.DOTALL | re.IGNORECASE | re.MULTILINE)
            if matches and len(matches) > 0:
                content = matches[0]
                # Handle if content is a tuple
                if isinstance(content, tuple) and len(content) > 0:
                    content = content[0]
        
        # Ensure content is a string
        if not isinstance(content, str):
            if isinstance(content, (list, tuple)) and len(content) > 0:
                content = content[0]
            else:
                content = str(content)
        
        content = content.strip()
        
        # Add the section to the document
        if content:
            # Add section heading - Remove the asterisks
            heading = doc.add_heading(level=2)
            heading_run = heading.add_run(formatted_section_title)
            heading_run.bold = True
            # Set heading text to black
            heading_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add section content with proper formatting including bullets
            add_content_with_bullets(doc, content)
    
    # Add horizontal line
    doc.add_paragraph("----------------------------------------------------------------------------")
    
    # Add AI Maturity Levels section (footer) - Remove the asterisks
    maturity_heading = doc.add_heading(level=2)
    maturity_heading_run = maturity_heading.add_run("AI Maturity Levels")
    maturity_heading_run.bold = True
    # Set heading text to black
    maturity_heading_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Low maturity level - with smaller font
    p = doc.add_paragraph()
    low_label = p.add_run("Low: ")
    low_label.bold = True
    low_label.font.size = Pt(9)  # Smaller font size
    low_text = p.add_run("Companies that are in the early stages of AI integration or development and/or typically in the ideation phase and/or with only a proof of concept. They have limited data, resources, and expertise, and a minimal understanding of AI. AI is minimally or not at all used in workflows, with no data management processes or AI roadmap in place.")
    low_text.font.size = Pt(9)  # Smaller font size
    
    # Moderate maturity level - with smaller font
    p = doc.add_paragraph()
    moderate_label = p.add_run("Moderate: ")
    moderate_label.bold = True
    moderate_label.font.size = Pt(9)  # Smaller font size
    moderate_text = p.add_run("Companies that are progressing in their AI journey, moving beyond the proof of concept stage with functional solutions. They have adequate data, resources, expertise, and understanding of AI. AI is either fully or partially integrated into their workflows, supported by established or developing data management processes, and guided by a partially or fully formulated AI roadmap.")
    moderate_text.font.size = Pt(9)  # Smaller font size
    
    # High maturity level - with smaller font
    p = doc.add_paragraph()
    high_label = p.add_run("High: ")
    high_label.bold = True
    high_label.font.size = Pt(9)  # Smaller font size
    high_text = p.add_run("Companies that have already developed advanced AI products and have an established customer base. AI is fully or partially integrated into their workflows, supported by established data management processes, and guided by an AI roadmap. They require assistance with specific technical details or when developing new AI applications on top of their existing solutions.")
    high_text.font.size = Pt(9)  # Smaller font size
    
    # Save the document
    doc.save(file_path)
    
    return file_path


